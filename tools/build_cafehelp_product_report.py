from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(
    r"C:\Users\shakur\Desktop\Документация_CafeHelp\Docs"
    r"\CafeHelp_Продуктовый_отчет_и_коммерческая_модель_2026.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "19324A"
TEAL = "2F8C7A"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F4F1"
PALE_GRAY = "F2F4F7"
MID_GRAY = "64748B"
DARK = "1F2937"
WHITE = "FFFFFF"
LINE = "CBD5E1"
AMBER = "B7791F"
PALE_AMBER = "FFF7E6"
GREEN = "2F855A"
RED = "B42318"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_width(table, width_dxa=9360, indent_dxa=0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep(paragraph, keep_next=False, keep_lines=True, page_break_before=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    if keep_lines:
        node = OxmlElement("w:keepLines")
        p_pr.append(node)
    if page_break_before:
        node = OxmlElement("w:pageBreakBefore")
        p_pr.append(node)


def set_repeat_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(end)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(run_color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_paragraph_border(paragraph, side: str, color: str, size=12, space=8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_run_font(run, size=None, bold=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_cell(cell) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)


def add_cell_text(cell, text, *, bold=False, color=DARK, size=9.5, align=None) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 14, MID_GRAY, 0, 10),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.widow_control = True

    if "Small Note" not in styles:
        small_note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small_note = styles["Small Note"]
    small_note.font.name = "Calibri"
    small_note._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    small_note.font.size = Pt(8.5)
    small_note.font.color.rgb = RGBColor.from_string(MID_GRAY)
    small_note.paragraph_format.space_after = Pt(4)
    small_note.paragraph_format.line_spacing = 1.0

    if "KPI Value" not in styles:
        kpi = styles.add_style("KPI Value", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kpi = styles["KPI Value"]
    kpi.font.name = "Calibri"
    kpi._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    kpi.font.size = Pt(18)
    kpi.font.bold = True
    kpi.font.color.rgb = RGBColor.from_string(NAVY)
    kpi.paragraph_format.space_after = Pt(2)

    core = doc.core_properties
    core.title = "CafeHelp — продуктовый отчёт и коммерческая модель"
    core.subject = "Описание продукта, развитие, функциональные улучшения, стоимость и требования к инфраструктуре"
    core.author = "Проект CafeHelp"
    core.keywords = "CafeHelp, автоматизация общепита, POS, склад, кухня, коммерческая модель"
    core.comments = "Подготовлено на основании фактического состава проекта CafeHelp и открытых рыночных данных."
    set_repeat_on_open(doc)


def build_header_footer(section) -> None:
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    widths = [4000, 5360]
    for idx, width in enumerate(widths):
        set_cell_width(table.cell(0, idx), width)
        set_cell_margins(table.cell(0, idx), top=0, bottom=50, start=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    r = left.add_run("CAFEHELP")
    set_run_font(r, size=10, bold=True, color=NAVY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    r = right.add_run("ПРОДУКТОВЫЙ ОТЧЁТ • 2026")
    set_run_font(r, size=8, bold=True, color=MID_GRAY)
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 8, "space": 0, "color": BLUE})

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    set_paragraph_border(p, "top", LINE, size=4, space=4)
    r = p.add_run("CafeHelp  •  Конфиденциальный рабочий документ  •  ")
    set_run_font(r, size=8, color=MID_GRAY)
    add_field(p, "PAGE", "1")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CAFEHELP")
    set_run_font(r, size=13, bold=True, color=TEAL)
    set_paragraph_border(p, "bottom", TEAL, size=20, space=8)

    doc.add_paragraph("", style="Normal")
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(38)
    p.add_run("Продуктовый отчёт\nи коммерческая модель")

    p = doc.add_paragraph(style="Subtitle")
    p.add_run(
        "Единая система управления кафе и рестораном:\n"
        "от кассы и кухни до склада, персонала и управленческой аналитики"
    )

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for i, width in enumerate((3120, 3120, 3120)):
        cell = table.cell(0, i)
        set_cell_width(cell, width)
        set_cell_margins(cell, top=160, bottom=160, start=160, end=160)
        set_cell_shading(cell, (PALE_BLUE, PALE_TEAL, PALE_GRAY)[i])
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": LINE},
            bottom={"val": "single", "sz": 4, "color": LINE},
            start={"val": "single", "sz": 4, "color": LINE},
            end={"val": "single", "sz": 4, "color": LINE},
        )
    cover_items = (
        ("НАЗНАЧЕНИЕ", "Управление операционной деятельностью заведения"),
        ("СТАТУС", "Рабочая система, готовая к пилотной эксплуатации"),
        ("МОДЕЛЬ", "Облако по подписке или локальная лицензия"),
    )
    for idx, (label, text) in enumerate(cover_items):
        cell = table.cell(0, idx)
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(label)
        set_run_font(r, size=8, bold=True, color=BLUE)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Подготовлено для принятия продуктовых и коммерческих решений")
    set_run_font(r, size=11, bold=True, color=NAVY)
    p = doc.add_paragraph()
    r = p.add_run("Версия 1.0  •  28 июля 2026 года  •  Москва")
    set_run_font(r, size=9.5, color=MID_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_border(p, "top", LINE, size=4, space=8)
    r = p.add_run(
        "Документ отражает текущее состояние проекта и рекомендуемую модель развития. "
        "Цены являются ориентировочными и подлежат уточнению после пилотного внедрения."
    )
    set_run_font(r, size=8.5, color=MID_GRAY)
    doc.add_page_break()


def add_section_title(doc, number: str, title: str, lead: str | None = None) -> None:
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(f"{number}  ")
    set_run_font(r, size=16, bold=True, color=TEAL)
    p.add_run(title)
    if lead:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(lead)
        set_run_font(r, size=11, color=MID_GRAY)


def add_callout(doc, title: str, text: str, fill=PALE_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_margins(cell, top=150, bottom=150, start=200, end=200)
    set_cell_shading(cell, fill)
    set_cell_border(cell, start={"val": "single", "sz": 22, "color": accent})
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper())
    set_run_font(r, size=8.5, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kpi_row(doc) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    data = (
        ("1 контур", "касса • кухня • склад"),
        ("2 роли", "владелец • сотрудник"),
        ("2 модели", "подписка • локально"),
    )
    for i, (value, label) in enumerate(data):
        cell = table.cell(0, i)
        set_cell_width(cell, 3120)
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        set_cell_shading(cell, WHITE)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            start={"val": "single", "sz": 6, "color": LINE},
            end={"val": "single", "sz": 6, "color": LINE},
        )
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.style = doc.styles["KPI Value"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(value)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.5, color=MID_GRAY)


def add_summary(doc: Document) -> None:
    add_section_title(
        doc,
        "01",
        "Резюме",
        "Краткое представление продукта, его ценности и рекомендуемого позиционирования.",
    )
    add_callout(
        doc,
        "Ключевой вывод",
        "CafeHelp — единая веб-система для ежедневного управления заведением общественного питания. "
        "Продукт объединяет продажи, кухню, склад, закупки, технологические карты, клиентов, персонал "
        "и аналитику, уменьшая количество разрозненных таблиц и ручных операций.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_kpi_row(doc)

    doc.add_heading("Что это за продукт", level=2)
    doc.add_paragraph(
        "CafeHelp предназначен для независимых кафе, кофеен, небольших ресторанов и локальных сетей. "
        "Система поддерживает полный операционный цикл: от поступления продуктов и расчёта себестоимости "
        "до оформления заказа, передачи его на кухню, закрытия смены и анализа результатов."
    )
    doc.add_paragraph(
        "Интерфейс работает в браузере и адаптирован под рабочие места владельца, кассира и кухни. "
        "Развёртывание возможно как на сервере поставщика, так и внутри инфраструктуры заказчика."
    )

    doc.add_heading("Позиционирование", level=2)
    for text in (
        "Практичный центр управления заведением — без необходимости собирать кассу, склад и аналитику из отдельных сервисов.",
        "Основная целевая аудитория — заведения с одной–тремя точками, которым важны прозрачность учёта и быстрый запуск.",
        "Текущий статус — рабочая система, целесообразная для контролируемого пилота; коммерческий масштаб следует наращивать после проверки в реальном заведении.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Ценность для бизнеса", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    headers = ("Для владельца", "Для персонала", "Для гостя")
    widths = (3120, 3120, 3120)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.5)
    row = table.add_row()
    values = (
        "Контроль выручки, себестоимости, остатков, долгов и работы смены в одном месте.",
        "Понятные рабочие сценарии: принять заказ, приготовить, списать, переместить, провести инвентаризацию.",
        "Более предсказуемое обслуживание, контроль готовности заказа и сохранение истории взаимодействия.",
    )
    for i, value in enumerate(values):
        cell = row.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        set_cell_shading(cell, (PALE_BLUE, PALE_GRAY, PALE_TEAL)[i])
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": LINE},
            bottom={"val": "single", "sz": 4, "color": LINE},
            start={"val": "single", "sz": 4, "color": LINE},
            end={"val": "single", "sz": 4, "color": LINE},
        )
        add_cell_text(cell, value, size=9.2)
    prevent_row_split(row)


def add_current_features(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "02",
        "Что реализовано",
        "Функциональные возможности, подтверждённые текущей структурой приложения.",
    )
    features = [
        (
            "Продажи и касса",
            "Открытие и закрытие смен, формирование заказа, добавление блюд и наборов, оплата, выдача, "
            "работа с доставкой и долгами, Z-отчёт, печать документов средствами браузера.",
        ),
        (
            "Кухонный экран",
            "Очередь заказов по текущей смене, статусы приготовления, контроль времени и передача готовности на кассу.",
        ),
        (
            "Меню и технологические карты",
            "Блюда, категории, наборы, изображения, состав и расход ингредиентов, отходы, единицы измерения, "
            "расчёт себестоимости блюд и заготовок.",
        ),
        (
            "Склад и производство",
            "Остатки продуктов и заготовок, поступления, списания, перемещения, производство полуфабрикатов, "
            "инвентаризация, оборот и отчёт по смене.",
        ),
        (
            "Закупки и поставщики",
            "Справочник поставщиков, ассортимент и цены, избранные позиции, приходные накладные и оприходование.",
        ),
        (
            "Клиенты и взаиморасчёты",
            "Карточка клиента, история и избранные заказы, учёт задолженности и сроков погашения, привязка VK.",
        ),
        (
            "Персонал",
            "Сотрудники и роли, учёт отработанных дней, начисления, выплаты и архив операций.",
        ),
        (
            "Управленческая панель",
            "Ключевые показатели, быстрые действия, популярные блюда, предупреждения по остаткам и аналитические отчёты.",
        ),
        (
            "Интеграционный контур",
            "Передача оплаченных заказов во внешний налоговый/партнёрский контур с очередью повторных отправок; "
            "VK-бот для связи клиента с профилем, просмотра истории и задолженности.",
        ),
        (
            "Прогнозирование и рекомендации",
            "Экспериментальный модуль прогноза спроса, анализа ингредиентов и генерации рекомендаций. "
            "Результаты следует использовать как поддержку решения, а не как автоматическое распоряжение.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths = (2500, 6860)
    for i, header in enumerate(("Функциональный блок", "Что получает пользователь")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.5)
    set_repeat_table_header(table.rows[0])
    for idx, (name, description) in enumerate(features):
        row = table.add_row()
        for i, value in enumerate((name, description)):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=85, bottom=85, start=120, end=120)
            set_cell_shading(cell, PALE_GRAY if idx % 2 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, value, bold=(i == 0), color=NAVY if i == 0 else DARK, size=8.8)
        prevent_row_split(row)

    doc.add_heading("Ролевой доступ", level=2)
    doc.add_paragraph(
        "Базовое разделение ролей предусматривает владельца и сотрудника. Владелец управляет справочниками, "
        "закупками, складом, персоналом и аналитикой; сотрудник работает в разрешённых операционных разделах. "
        "Для коммерческой эксплуатации матрицу прав рекомендуется детализировать по должностям."
    )
    add_callout(
        doc,
        "Граница продукта",
        "CafeHelp не заменяет фискальный регистратор, ОФД, банковский эквайринг или официального оператора "
        "налогового обмена. Эти сервисы и оборудование подключаются отдельно в соответствии с законодательством "
        "и договором заказчика.",
        fill=PALE_AMBER,
        accent=AMBER,
    )


def add_development_paths(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "03",
        "Пути развития",
        "Стратегические направления, которые повышают рыночную ценность продукта.",
    )
    directions = [
        (
            "01",
            "Продукт для одной точки",
            "Сделать запуск максимально быстрым: типовые настройки, готовые роли, импорт меню и понятный сценарий первого рабочего дня.",
            "Снижение стоимости внедрения и барьера перехода с таблиц.",
        ),
        (
            "02",
            "Управление сетью",
            "Единый кабинет владельца, сравнение заведений, централизованные меню и цены, перенос лучших практик между точками.",
            "Рост среднего чека продукта и удержание клиентов.",
        ),
        (
            "03",
            "Экосистема гостя",
            "Объединить профиль гостя, лояльность, онлайн-заказ, бронирование, обратную связь и персональные предложения.",
            "Повторные продажи и рост частоты посещений.",
        ),
        (
            "04",
            "Закупки и контроль маржи",
            "Сравнение цен поставщиков, планирование закупок, контроль отклонений себестоимости и согласование крупных закупок.",
            "Снижение потерь и защита маржинальности.",
        ),
        (
            "05",
            "Помощник руководителя",
            "Перевести аналитику из набора графиков в понятные действия: что заказать, какое блюдо пересчитать, где растёт списание.",
            "Быстрые решения без отдельного аналитика.",
        ),
        (
            "06",
            "Открытая интеграционная платформа",
            "Подключение бухгалтерии, доставки, платежей, программ лояльности и отраслевых сервисов через каталог интеграций.",
            "Меньше ручного переноса данных и выше ценность экосистемы.",
        ),
    ]
    for number, title, action, result in directions:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table)
        left, right = table.rows[0].cells
        set_cell_width(left, 940)
        set_cell_width(right, 8420)
        for cell in (left, right):
            set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
        set_cell_shading(left, TEAL)
        set_cell_shading(right, WHITE)
        add_cell_text(left, number, bold=True, color=WHITE, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_cell(right)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        set_run_font(r, size=11, bold=True, color=NAVY)
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(action)
        set_run_font(r, size=9.3, color=DARK)
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Бизнес-эффект: ")
        set_run_font(r, size=8.8, bold=True, color=TEAL)
        r = p.add_run(result)
        set_run_font(r, size=8.8, color=MID_GRAY)
        prevent_row_split(table.rows[0])
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)


def add_feature_roadmap(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "04",
        "Рекомендуемые функциональные улучшения",
        "Новые пользовательские возможности без углубления в техническую реализацию.",
    )
    roadmap = [
        (
            "0–6 месяцев",
            "Пилот и ежедневная полезность",
            [
                "Программа лояльности: бонусы, уровни, дни рождения и персональные предложения.",
                "QR-меню с оформлением заказа за столом или на самовывоз.",
                "Бронирование столов и календарь загрузки зала.",
                "Стоп-лист, автоматически доступный кассе, кухне и гостевому меню.",
                "Импорт меню и остатков из шаблона для быстрого старта.",
                "Мобильная сводка владельца с ключевыми событиями дня.",
            ],
        ),
        (
            "6–12 месяцев",
            "Рост выручки и контроль команды",
            [
                "Онлайн-заказ и собственная страница доставки без комиссии агрегатора.",
                "Интеграции с агрегаторами доставки и единая очередь заказов.",
                "Отзывы гостей, NPS и обработка претензий внутри карточки клиента.",
                "Расписание смен, замены сотрудников и уведомления о выходе.",
                "Согласование закупок и сравнение предложений поставщиков.",
                "Маркетинговые сегменты: новые, постоянные, засыпающие и VIP-гости.",
            ],
        ),
        (
            "12–24 месяца",
            "Сеть и тиражирование",
            [
                "Единый кабинет сети с рейтингом точек и консолидированной отчётностью.",
                "Централизованное меню с локальными ценами и ассортиментом.",
                "Франчайзинговые стандарты, чек-листы и контроль исполнения.",
                "Подарочные сертификаты, абонементы и корпоративное питание.",
                "Личный кабинет гостя и брендированное мобильное приложение.",
                "Каталог готовых интеграций с бухгалтерией, платежами и доставкой.",
            ],
        ),
    ]
    for period, focus, items in roadmap:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table)
        left, right = table.rows[0].cells
        set_cell_width(left, 2300)
        set_cell_width(right, 7060)
        for cell in (left, right):
            set_cell_margins(cell, top=140, bottom=140, start=170, end=170)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
        set_cell_shading(left, NAVY)
        set_cell_shading(right, PALE_GRAY)
        clear_cell(left)
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(period.upper())
        set_run_font(r, size=8.5, bold=True, color=TEAL)
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(focus)
        set_run_font(r, size=11, bold=True, color=WHITE)
        clear_cell(right)
        for idx, item in enumerate(items):
            p = right.paragraphs[0] if idx == 0 else right.add_paragraph()
            p.style = doc.styles["List Bullet"]
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(item)
            set_run_font(r, size=9.1, color=DARK)
        prevent_row_split(table.rows[0])
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_callout(
        doc,
        "Принцип приоритизации",
        "В первую очередь следует развивать функции, которые ежедневно видит клиент и которые можно измерить "
        "через выручку, скорость обслуживания, повторные визиты или сокращение потерь. Сложные прогнозные функции "
        "целесообразно расширять после накопления качественных данных пилотных заведений.",
        fill=PALE_BLUE,
        accent=BLUE,
    )


def add_pricing(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "05",
        "Стоимость и коммерческая модель",
        "Рекомендуемые цены для проверки спроса. Все суммы указаны в рублях и не включают налоги, оборудование и услуги сторонних операторов.",
    )
    add_callout(
        doc,
        "Рекомендация",
        "Основной моделью сделать подписку: она снижает входной барьер и финансирует поддержку продукта. "
        "Локальную бессрочную лицензию оставить для организаций с требованиями к размещению данных внутри своей инфраструктуры.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    doc.add_heading("Облачная подписка", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths = (2000, 1750, 1750, 3860)
    headers = ("Тариф", "В месяц", "За год", "Состав и ограничения")
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.2)
    set_repeat_table_header(table.rows[0])
    plans = [
        (
            "Старт",
            "3 900 ₽",
            "39 000 ₽",
            "1 заведение, 1 касса, до 5 пользователей: меню, смены, продажи, базовый склад и стандартные отчёты.",
        ),
        (
            "Бизнес",
            "6 900 ₽",
            "69 000 ₽",
            "До 3 касс и 15 пользователей: полный склад, техкарты, заготовки, поставщики, кухня, клиенты, долги и VK.",
        ),
        (
            "Про",
            "9 900 ₽",
            "99 000 ₽",
            "До 5 касс и 50 пользователей: всё из «Бизнес», расширенная аналитика, прогнозы, интеграционный контур и приоритетная поддержка.",
        ),
    ]
    for idx, plan in enumerate(plans):
        row = table.add_row()
        for i, value in enumerate(plan):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=105, bottom=105, start=120, end=120)
            set_cell_shading(cell, PALE_TEAL if idx == 1 else (PALE_GRAY if idx % 2 else WHITE))
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(
                cell,
                value,
                bold=(i in (0, 1, 2)),
                color=TEAL if idx == 1 and i in (0, 1, 2) else (NAVY if i in (0, 1, 2) else DARK),
                size=9.1,
            )
        prevent_row_split(row)

    p = doc.add_paragraph(style="Small Note")
    p.add_run(
        "Годовая оплата соответствует примерно десяти месяцам использования. Дополнительное заведение — 70% "
        "стоимости выбранного тарифа. Тестовый период — 30 дней после первичной настройки."
    )

    doc.add_heading("Внедрение и локальная лицензия", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths2 = (3300, 1900, 4160)
    for i, header in enumerate(("Позиция", "Цена", "Комментарий")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths2[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.2)
    items = [
        ("Первичная настройка «Старт»", "15 000 ₽", "Импорт базовых справочников и удалённое обучение."),
        ("Внедрение «Бизнес»", "35 000 ₽", "Настройка процессов, перенос меню и обучение ключевых пользователей."),
        ("Внедрение «Про»", "60 000 ₽", "Расширенная настройка, интеграции и сопровождение пилотного запуска."),
        ("Бессрочная локальная лицензия", "280 000 ₽", "Одно заведение, полный функциональный состав без оборудования."),
        ("Развёртывание локальной версии", "60 000 ₽", "Установка, первичная настройка, резервирование и обучение."),
        ("Поддержка локальной версии", "56 000 ₽/год", "Обновления и консультации; 20% стоимости лицензии."),
        ("Дополнительное заведение локально", "140 000 ₽", "Лицензия на следующую точку в инфраструктуре заказчика."),
    ]
    for idx, item in enumerate(items):
        row = table.add_row()
        for i, value in enumerate(item):
            cell = row.cells[i]
            set_cell_width(cell, widths2[i])
            set_cell_margins(cell, top=85, bottom=85, start=120, end=120)
            set_cell_shading(cell, PALE_GRAY if idx % 2 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, value, bold=(i == 1), color=NAVY if i in (0, 1) else DARK, size=8.8)
        prevent_row_split(row)

    doc.add_heading("Что не входит в стоимость", level=2)
    for text in (
        "кассовое и кухонное оборудование, фискальный регистратор, принтеры, сканеры, планшеты и сервер;",
        "ОФД, эквайринг, электронная подпись, оператор фискальных данных и услуги официального налогового партнёра;",
        "платные сторонние интеграции, SMS/мессенджер-рассылки и индивидуальная доработка под процессы заказчика;",
        "выездные работы, нестандартный перенос данных и обучение сверх согласованного пакета.",
    ):
        doc.add_paragraph(text, style="List Bullet")


def add_market_context(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "06",
        "Обоснование цены",
        "Сопоставление с публичными предложениями российского рынка по состоянию на 28 июля 2026 года.",
    )
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths = (1900, 2700, 4760)
    for i, header in enumerate(("Решение", "Публичный ориентир", "Комментарий")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.2)
    market = [
        (
            "Restik",
            "2 499 ₽/мес. за год\n3 490 ₽ помесячно",
            "Единая подписка с учётом, кассой, лояльностью, сайтом доставки и QR-меню.",
        ),
        (
            "Quick Resto",
            "4 990 / 7 990 /\n9 990 ₽ в месяц",
            "Тарифы «Старт», «Про» и «Энтерпрайз»; годовая оплата дешевле помесячной.",
        ),
        (
            "Saby Presto",
            "24 000 / 37 000 /\n49 000 ₽ за год",
            "Публичные годовые уровни для кассы, обслуживания у столов и крупных заведений.",
        ),
        (
            "iiko",
            "Модульная модель",
            "Отдельно тарифицируются облачные сервисы, кухонные экраны, доставка, лояльность и интеграции.",
        ),
    ]
    for idx, item in enumerate(market):
        row = table.add_row()
        for i, value in enumerate(item):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=95, bottom=95, start=120, end=120)
            set_cell_shading(cell, PALE_GRAY if idx % 2 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, value, bold=(i in (0, 1)), color=NAVY if i in (0, 1) else DARK, size=8.9)
        prevent_row_split(row)

    doc.add_heading("Логика позиционирования CafeHelp", level=2)
    for text in (
        "«Старт» находится между минимальными едиными подписками и полноценными рыночными тарифами — это облегчает вход в пилот.",
        "«Бизнес» соответствует основному сценарию CafeHelp и монетизирует сильные стороны продукта: склад, техкарты, закупки, кухню и клиентов.",
        "«Про» остаётся ниже или на уровне верхних публичных тарифов, но включает прогнозирование и интеграционный контур; продавать его следует только после подтверждения стабильности этих функций.",
        "Локальная лицензия окупается для заказчика примерно за 3–4 года по сравнению с тарифом «Бизнес», при этом расходы на собственный сервер и администрирование остаются на стороне заказчика.",
    ):
        doc.add_paragraph(text, style="List Number")

    add_callout(
        doc,
        "Коммерческий эксперимент",
        "Первые 5–10 внедрений рекомендуется проводить по пилотной цене со скидкой 30–40% в обмен на регулярную "
        "обратную связь, право использовать обезличенные метрики и согласованный публичный кейс. После пилотов "
        "тарифы пересматриваются по фактической нагрузке на поддержку и ценности для клиента.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    doc.add_heading("Источники рыночных ориентиров", level=2)
    sources = [
        ("Quick Resto — тарифы", "https://quickresto.ru/price/"),
        ("Restik — автоматизация кафе", "https://restik.com/automation/kafe/"),
        ("Saby Presto — тарифы для общепита", "https://saby.ru/presto"),
        ("iiko — цены и состав решений", "https://iiko.ru/solutions/products/price"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="Small Note")
        add_hyperlink(p, label, url)
        r = p.add_run(f" — {url}")
        set_run_font(r, size=8.5, color=MID_GRAY)
    p = doc.add_paragraph(style="Small Note")
    p.add_run(
        "Публичные цены могут изменяться и не учитывают персональные скидки, оборудование, внедрение и платные дополнения. "
        "Сопоставление используется только как ориентир для продуктового решения."
    )


def add_local_requirements(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "07",
        "Требования при локальном развёртывании",
        "Конфигурации для размещения CafeHelp непосредственно в заведении или офисе заказчика.",
    )
    doc.add_heading("Сервер внутри заведения", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths = (2400, 3420, 3540)
    for i, header in enumerate(("Параметр", "Минимум для пилота", "Рекомендуется для работы")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.1)
    rows = [
        ("Сценарий", "1 точка, до 5 одновременных пользователей", "1 точка, 5–20 одновременных пользователей"),
        ("Процессор", "x86-64, 4 физических/логических ядра", "x86-64, 8 ядер"),
        ("Оперативная память", "16 ГБ", "32 ГБ"),
        ("Системный диск", "SSD 150 ГБ", "NVMe SSD 500 ГБ"),
        ("Надёжность дисков", "Один SSD + ежедневная внешняя копия", "2 × 500 ГБ SSD/NVMe в RAID 1"),
        ("Сеть", "1 Гбит/с внутри заведения; интернет от 20 Мбит/с", "1 Гбит/с; резервный интернет-канал"),
        ("Электропитание", "ИБП от 600 ВА", "ИБП 1 000 ВА с корректным завершением работы"),
        ("Резервное хранение", "Внешний диск от 250 ГБ", "NAS/внешнее хранилище от 1 ТБ"),
        ("Операционная система", "Ubuntu 22.04/24.04 LTS; Windows 11 Pro — для пилота", "Ubuntu 24.04 LTS"),
        ("Графический ускоритель", "Не требуется", "Не требуется"),
    ]
    for idx, item in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(item):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=82, bottom=82, start=120, end=120)
            set_cell_shading(cell, PALE_GRAY if idx % 2 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, value, bold=(i == 0), color=NAVY if i == 0 else DARK, size=8.6)
        prevent_row_split(row)

    doc.add_heading("Рабочие места", level=2)
    workstations = [
        (
            "Касса",
            "2 ядра, 4 ГБ ОЗУ, экран от 1280×720, актуальная версия Chrome/Edge; термопринтер 80 мм и сканер — по необходимости.",
        ),
        (
            "Кухня",
            "4 ГБ ОЗУ, экран 1920×1080 или планшет от 10″, защищённое от влаги размещение, стабильное Wi‑Fi/LAN-соединение.",
        ),
        (
            "Владелец/офис",
            "Обычный ноутбук или ПК с 8 ГБ ОЗУ и современным браузером; удалённый доступ — только через защищённое соединение.",
        ),
    ]
    for title, text in workstations:
        p = doc.add_paragraph(style="Heading 3")
        p.add_run(title)
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.2)

    add_callout(
        doc,
        "Критичное условие",
        "Локальный сервер нельзя размещать на кассовом компьютере, который ежедневно выключается. "
        "Он должен работать непрерывно, быть подключён к ИБП и выполнять автоматическое резервное копирование.",
        fill=PALE_AMBER,
        accent=AMBER,
    )


def add_server_requirements(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "08",
        "Требования к серверу",
        "Рекомендуемые параметры для облачного или выделенного размещения.",
    )
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    widths = (2050, 1960, 1960, 3390)
    for i, header in enumerate(("Сценарий", "CPU / RAM", "Диск", "Ориентир нагрузки")):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, header, bold=True, color=WHITE, size=9.1)
    server_plans = [
        ("Демо / пилот", "4 vCPU\n12 ГБ", "150 ГБ SSD", "1 заведение, до 5 одновременных пользователей"),
        ("Рабочая точка", "8 vCPU\n16 ГБ", "250 ГБ NVMe", "1–2 заведения, до 25 одновременных пользователей"),
        ("Небольшая сеть", "12 vCPU\n32 ГБ", "500 ГБ NVMe", "3–10 заведений, до 100 одновременных пользователей"),
    ]
    for idx, item in enumerate(server_plans):
        row = table.add_row()
        for i, value in enumerate(item):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
            set_cell_shading(cell, PALE_TEAL if idx == 1 else (PALE_GRAY if idx % 2 else WHITE))
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, value, bold=(i in (0, 1, 2)), color=NAVY if i in (0, 1, 2) else DARK, size=8.9)
        prevent_row_split(row)

    doc.add_heading("Обязательные эксплуатационные условия", level=2)
    conditions = [
        ("ОС и запуск", "Ubuntu 24.04 LTS, контейнерный запуск, автоматический рестарт сервисов."),
        ("Доступ", "Доменное имя, статический IP или стабильный DNS, HTTPS-сертификат, закрытые служебные порты."),
        ("Резервирование", "Ежедневная копия баз и файлов; хранение 30 ежедневных и 12 еженедельных копий вне основного сервера."),
        ("Мониторинг", "Контроль доступности, загрузки CPU/RAM, свободного места, ошибок приложения и успешности резервных копий."),
        ("Связь", "Стабильный интернет от 50 Мбит/с для серверного размещения; желательно два независимых канала."),
        ("Разделение данных", "База данных и пользовательские изображения должны резервироваться раздельно; доступ — по минимально необходимым правам."),
        ("Масштабирование", "При заполнении диска на 70% или устойчивой загрузке CPU выше 65% требуется плановое расширение."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for idx, (name, value) in enumerate(conditions):
        row = table.add_row()
        for i, text in enumerate((name, value)):
            cell = row.cells[i]
            set_cell_width(cell, (2300, 7060)[i])
            set_cell_margins(cell, top=92, bottom=92, start=130, end=130)
            set_cell_shading(cell, PALE_GRAY if idx % 2 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
            add_cell_text(cell, text, bold=(i == 0), color=NAVY if i == 0 else DARK, size=8.9)
        prevent_row_split(row)

    doc.add_heading("Целевые показатели эксплуатации", level=2)
    for text in (
        "доступность сервиса для рабочего тарифа — не ниже 99,5% в месяц;",
        "целевое время восстановления после сбоя — до 4 часов;",
        "допустимая потеря данных при штатной политике резервирования — не более 24 часов;",
        "проверка восстановления из резервной копии — не реже одного раза в квартал;",
        "графический ускоритель не требуется: текущие прогнозные расчёты выполняются на CPU.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    p = doc.add_paragraph(style="Small Note")
    p.add_run(
        "Нагрузочные ориентиры являются стартовыми. Итоговая конфигурация определяется после пилота с учётом "
        "числа заведений, касс, фотографий, срока хранения данных и интенсивности аналитических расчётов."
    )


def add_recommendations(doc: Document) -> None:
    doc.add_page_break()
    add_section_title(
        doc,
        "09",
        "Рекомендации по выводу на рынок",
        "Последовательность, которая снижает коммерческий риск и даёт измеримые результаты.",
    )
    steps = [
        (
            "1",
            "Провести пилот в одном заведении",
            "4–8 недель реальной работы кассы, кухни и склада. Зафиксировать исходные показатели и ответственного со стороны заведения.",
        ),
        (
            "2",
            "Подтвердить ценность",
            "Измерить скорость закрытия смены, точность остатков, время приготовления, количество ручных операций и использование отчётов владельцем.",
        ),
        (
            "3",
            "Упаковать внедрение",
            "Подготовить типовой импорт меню, регламент запуска, обучение по ролям и критерии готовности к первому рабочему дню.",
        ),
        (
            "4",
            "Продать первые подписки",
            "Предложить пилотным клиентам специальную цену и перейти к стандартным тарифам после подтверждения стабильности.",
        ),
        (
            "5",
            "Масштабировать через партнёров",
            "Подключить интеграторов кассового оборудования и отраслевых консультантов, сохранив единые стандарты качества.",
        ),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for idx, (number, title, text) in enumerate(steps):
        row = table.add_row()
        left, right = row.cells
        set_cell_width(left, 950)
        set_cell_width(right, 8410)
        for cell in (left, right):
            set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": LINE},
                bottom={"val": "single", "sz": 4, "color": LINE},
                start={"val": "single", "sz": 4, "color": LINE},
                end={"val": "single", "sz": 4, "color": LINE},
            )
        set_cell_shading(left, TEAL if idx < 3 else BLUE)
        set_cell_shading(right, PALE_GRAY if idx % 2 else WHITE)
        add_cell_text(left, number, bold=True, color=WHITE, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_cell(right)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.2, color=DARK)
        prevent_row_split(row)

    doc.add_heading("Решение, рекомендуемое к утверждению", level=2)
    add_callout(
        doc,
        "Итог",
        "Позиционировать CafeHelp как систему для независимых заведений и небольших сетей; запустить три облачных "
        "тарифа от 3 900 до 9 900 ₽ в месяц; сохранить локальную лицензию за 280 000 ₽; ближайшие продуктовые "
        "приоритеты — лояльность, QR-меню, бронирование, онлайн-заказ и мобильная сводка владельца.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    doc.add_heading("Основание документа", level=2)
    doc.add_paragraph(
        "Описание текущих возможностей подготовлено по исходному коду, маршрутам интерфейса, конфигурации запуска "
        "и рабочим отчётам проекта CafeHelp, доступным на 28 июля 2026 года. Коммерческие цены и инфраструктурные "
        "параметры являются рекомендацией для пилотного этапа и должны быть пересмотрены по результатам фактической эксплуатации."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    set_paragraph_border(p, "top", BLUE, size=10, space=8)
    r = p.add_run("CafeHelp")
    set_run_font(r, size=12, bold=True, color=NAVY)
    r = p.add_run("  •  продуктовый отчёт и коммерческая модель  •  версия 1.0")
    set_run_font(r, size=9, color=MID_GRAY)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    build_header_footer(doc.sections[0])
    add_cover(doc)
    add_summary(doc)
    add_current_features(doc)
    add_development_paths(doc)
    add_feature_roadmap(doc)
    add_pricing(doc)
    add_market_context(doc)
    add_local_requirements(doc)
    add_server_requirements(doc)
    add_recommendations(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
