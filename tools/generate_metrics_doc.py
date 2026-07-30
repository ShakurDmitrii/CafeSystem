from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path("docs") / "Задание_3_Мониторинг_CafeHelp.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def set_default_font(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x16, 0x33, 0x5B)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2.font.size = Pt(12)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x16, 0x33, 0x5B)
    else:
        run.font.size = Pt(13)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = p.add_run(bold_prefix)
        prefix.bold = True
        prefix.font.name = "Times New Roman"
        prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        suffix = p.add_run(text[len(bold_prefix):])
        suffix.font.name = "Times New Roman"
        suffix._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_width(hdr_cells[idx], widths_cm[idx])
        set_cell_shading(hdr_cells[idx], "D9EAF7")
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(10)

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = value
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_width(row_cells[idx], widths_cm[idx])
            for paragraph in row_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)

    doc.add_paragraph()


def add_dashboard_scheme(doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(
        "+----------------------------------------------------------------------------------+\n"
        "| CafeHelp Service Dashboard                                                       |\n"
        "+----------------------------------------------------------------------------------+\n"
        "| SLA / Availability | P95 API | Error Rate | Print Success | MTTR               |\n"
        "+----------------------------------------------------------------------------------+\n"
        "| Delayed Orders     | Critical Stock | Active Alerts | Incidents Today          |\n"
        "+----------------------------------------------------------------------------------+\n"
        "| Trend: latency/error (24h)        | Trend: delayed orders / print success       |\n"
        "+----------------------------------------------------------------------------------+\n"
        "| Alert feed: time, severity, metric, object, owner, status                        |\n"
        "+----------------------------------------------------------------------------------+\n"
        "| Shift context: active shift, orders today, current worker, printer status        |\n"
        "+----------------------------------------------------------------------------------+"
    )
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    add_title(
        doc,
        "Система метрик и мониторинга для сервиса CafeHelp",
        "Задание 3. Контроль качества ИТ-сервиса, интерпретация показателей и система реагирования",
    )

    add_paragraph(
        doc,
        "Цель данной работы состоит в разработке системы метрик и мониторинга для проекта CafeHelp — сервиса автоматизации кафе, который включает кассу, управление сменами, складской учёт, печать чеков и аналитические функции. Система мониторинга должна позволять отслеживать как техническое состояние сервиса, так и его влияние на операционную деятельность кафе."
    )

    add_heading(doc, "1. Цель мониторинга", level=1)
    add_paragraph(
        doc,
        "Основная цель мониторинга сервиса CafeHelp заключается в обеспечении стабильной работы критичных бизнес-функций: создания заказа, печати чека, открытия и закрытия смен, доступа к складским остаткам и своевременного реагирования на инциденты. В отличие от чисто технического подхода, в данной работе мониторинг рассматривается как инструмент управления качеством сервиса и предотвращения потерь выручки."
    )

    add_heading(doc, "2. Ключевые метрики сервиса", level=1)
    add_paragraph(
        doc,
        "Для сервиса CafeHelp были выбраны метрики, которые отражают как техническую стабильность, так и операционное качество обслуживания клиентов. Выбор метрик связан с реальными рисками проекта: остановкой кассы, потерей заказов, задержками на кухне и отсутствием ингредиентов."
    )

    metrics_rows = [
        [
            "Доступность критичных функций",
            "Доля успешных проверок сценариев order, shift, warehouse и print.",
            "Если сервис недоступен, кафе не может принимать и исполнять заказы.",
            "Переключение на резервный сценарий, эскалация техподдержке, проверка БД и API.",
        ],
        [
            "P95 времени отклика",
            "95-й перцентиль ответа для критичных API.",
            "Показывает скорость сервиса под нагрузкой и влияет на удобство работы кассира.",
            "Оптимизация запросов, индексов, сетевого взаимодействия и очередей.",
        ],
        [
            "Уровень ошибок",
            "Доля запросов с ошибкой 5xx или неуспешным бизнес-результатом.",
            "Ошибки напрямую влияют на потерю заказов, повторный ввод и простой персонала.",
            "Анализ логов, исправление дефектов, rollback или hotfix.",
        ],
        [
            "Успешность печати чеков",
            "Доля успешно завершённых заданий печати.",
            "Если чек не напечатан, заказ может не попасть на кухню и клиент не будет обслужен вовремя.",
            "Проверка принтера, драйвера, очереди печати и Python-модуля печати.",
        ],
        [
            "Доля заказов с задержкой",
            "Процент заказов, в которых timeDelay больше нуля.",
            "Метрика отражает качество исполнения заказа и влияние сервиса на клиентский опыт.",
            "Коррекция процесса смены, интерфейса и работы кухни.",
        ],
        [
            "Количество критичных остатков",
            "Число ингредиентов, остаток которых ниже минимального порога.",
            "Дефицит ингредиентов делает невозможным приготовление блюд и ломает меню.",
            "Перезаказ, блокировка блюд, пересмотр закупок и уведомление менеджеру.",
        ],
        [
            "MTTR",
            "Среднее время восстановления после инцидента.",
            "Показывает зрелость процесса реагирования и способность сервиса быстро возвращаться в рабочее состояние.",
            "Уточнение инструкций, ролей дежурных и сценариев восстановления.",
        ],
    ]
    add_table(
        doc,
        ["Метрика", "Описание", "Почему важна", "Управленческие решения"],
        metrics_rows,
        [4.5, 6.2, 7.6, 7.8],
    )

    add_heading(doc, "3. Методика измерения", level=1)
    add_paragraph(
        doc,
        "Для каждой метрики была определена воспроизводимая методика измерения: источник данных, способ расчёта, периодичность и ограничения. Это необходимо для того, чтобы показатели можно было использовать в регулярном контроле сервиса, а не только в разовом анализе."
    )

    method_rows = [
        [
            "Доступность",
            "Health-check, логи Spring Boot и Python API",
            "успешные проверки / все проверки * 100%",
            "1 раз в минуту",
            "Не фиксирует ухудшение UX при формально доступном API",
        ],
        [
            "P95 отклика",
            "Access-логи, middleware, APM",
            "P95 для POST /api/orders, GET /api/shifts, POST /print/order",
            "Каждые 5 минут",
            "Зависит от нагрузки и распределения пользовательских сценариев",
        ],
        [
            "Уровень ошибок",
            "HTTP-логи, exception-логи",
            "ошибочные запросы / все запросы * 100%",
            "Каждые 5 минут",
            "Нужно разделять системные и пользовательские ошибки",
        ],
        [
            "Успешность печати",
            "Логи Python-модуля печати и статусы print jobs",
            "успешные задания печати / все задания печати * 100%",
            "Каждые 5 минут",
            "USB-принтер зависит и от драйвера, и от состояния ОС",
        ],
        [
            "Доля задержек",
            "Таблица заказов, поле timeDelay",
            "заказы с delay > 0 / все заказы * 100%",
            "По смене и по дню",
            "На задержку влияет не только ИТ, но и работа кухни",
        ],
        [
            "Критичные остатки",
            "Остатки по складу и складские движения",
            "count(product where qty <= min_threshold)",
            "Каждые 15 минут",
            "Порог зависит от единицы измерения и характера ингредиента",
        ],
        [
            "MTTR",
            "Журнал инцидентов, время ack и resolve",
            "avg(время закрытия - время открытия инцидента)",
            "По неделе и месяцу",
            "Требуется дисциплина фиксации инцидентов",
        ],
    ]
    add_table(
        doc,
        ["Метрика", "Источник данных", "Способ расчёта", "Периодичность", "Ограничения и допущения"],
        method_rows,
        [3.8, 5.6, 6.8, 3.2, 7.3],
    )

    add_heading(doc, "4. Макет дашборда", level=1)
    add_paragraph(
        doc,
        "Дашборд должен быть ориентирован на три группы пользователей: операторов смены, технического администратора и менеджмент. Верхняя часть показывает текущее состояние критичных функций, центральная часть — активные проблемы и тренды, нижняя часть — контекст смены и историю алертов."
    )
    add_paragraph(doc, "Структура дашборда представлена следующей схемой:")
    add_dashboard_scheme(doc)
    add_paragraph(
        doc,
        "На дашборде используются KPI-виджеты, графики трендов, лента активных алертов и контекстный блок текущей смены. Такое расположение позволяет одновременно оценивать общее состояние сервиса и быстро переходить к разбору конкретной проблемы."
    )

    add_heading(doc, "5. Пороги и алерты", level=1)
    add_paragraph(
        doc,
        "Для каждой метрики установлены три зоны контроля: нормальная, предупреждающая и критическая. Алерт формируется при пересечении порога в двух подряд окнах измерения, чтобы уменьшить количество ложных срабатываний. Критические уведомления отправляются техническому администратору и владельцу сервиса, а операционные — менеджеру смены."
    )

    alert_rows = [
        [
            "Доступность",
            ">= 99.5%",
            "< 99.5%",
            "< 99.0%",
            "Админ, владелец",
            "Проверка API, БД, сети, переход на резервный сценарий",
        ],
        [
            "P95 отклика",
            "<= 700 мс",
            "700-1500 мс",
            "> 1500 мс",
            "Техподдержка",
            "Анализ медленных запросов, нагрузки и очередей",
        ],
        [
            "Уровень ошибок",
            "< 1%",
            "1-3%",
            "> 3%",
            "Техподдержка, разработчик",
            "Поиск дефекта, rollback, hotfix",
        ],
        [
            "Успешность печати",
            ">= 98%",
            "95-98%",
            "< 95%",
            "Оператор, админ",
            "Проверка принтера, драйвера, очереди печати",
        ],
        [
            "Доля задержек",
            "< 10%",
            "10-20%",
            "> 20%",
            "Менеджер смены",
            "Проверка загрузки кухни и маршрута заказа",
        ],
        [
            "Критичные остатки",
            "0-2 SKU",
            "3-5 SKU",
            "> 5 SKU",
            "Кладовщик, менеджер",
            "Срочный дозакуп, временная блокировка блюд",
        ],
        [
            "MTTR",
            "<= 15 мин",
            "15-30 мин",
            "> 30 мин",
            "Руководитель ИТ/проекта",
            "Пересмотр регламента и процесса реагирования",
        ],
    ]
    add_table(
        doc,
        ["Метрика", "Норма", "Предупреждение", "Критично", "Получатель", "Действие"],
        alert_rows,
        [3.6, 2.5, 3.1, 2.7, 4.2, 10.0],
    )

    add_heading(doc, "6. Ценность выбранных метрик и алертов", level=1)
    add_paragraph(
        doc,
        "Выбранные метрики позволяют контролировать не только техническую исправность компонентов, но и фактическое качество сервиса для кафе. Доступность, отклик, ошибки и MTTR защищают SLA и помогают снижать длительность простоев. Успешность печати чеков предотвращает потерю заказов на кухне. Доля задержек отражает влияние сервиса на клиентский опыт. Контроль критичных остатков уменьшает риск отмены блюд и сбоев в работе смены."
    )
    add_bullets(
        doc,
        [
            "Снижение риска остановки продаж из-за недоступности кассы или смен.",
            "Снижение риска потери заказов из-за ошибок API и сбоев печати.",
            "Снижение вероятности ухудшения клиентского опыта из-за задержек.",
            "Предотвращение дефицита ингредиентов и срыва меню.",
            "Поддержание SLA по доступности сервиса и времени реакции на инциденты.",
        ],
    )

    add_heading(doc, "7. Заключение", level=1)
    add_paragraph(
        doc,
        "Для проекта CafeHelp система мониторинга должна быть построена вокруг критических пользовательских сценариев, а не только вокруг загрузки сервера и технических ресурсов. Предложенный набор метрик объединяет контроль доступности, скорости, ошибок, печати, исполнения заказов и состояния склада. Такой подход делает мониторинг инструментом управления качеством ИТ-сервиса и позволяет принимать обоснованные управленческие решения."
    )

    add_heading(doc, "Приложение А. Структура презентации", level=1)
    add_bullets(
        doc,
        [
            "Слайд 1. О проекте CafeHelp и его критичных функциях.",
            "Слайд 2. Почему сервису необходим мониторинг.",
            "Слайд 3. Выбранные метрики и логика их использования.",
            "Слайд 4. Методика измерения и источники данных.",
            "Слайд 5. Макет дашборда.",
            "Слайд 6. Пороги, алерты и сценарии реагирования.",
            "Слайд 7. Бизнес-ценность и влияние на SLA.",
        ],
    )

    add_heading(doc, "Приложение Б. Краткий питч", level=1)
    add_paragraph(
        doc,
        "Я разработал систему мониторинга для сервиса CafeHelp, который автоматизирует работу кафе: приём заказов, смены, склад и печать чеков. В основу мониторинга я заложил не только технические метрики, такие как доступность, отклик, ошибки и MTTR, но и операционные показатели: успешность печати, долю задержанных заказов и критичные остатки на складе. Это важно, потому что для такого сервиса даже локальный сбой быстро превращается в потерю выручки и ухудшение клиентского опыта. Я также определил пороговые значения, структуру дашборда и механизм алертов. В результате мониторинг позволяет заранее выявлять риски, быстрее восстанавливать сервис и поддерживать целевой SLA."
    )

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
