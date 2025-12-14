import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_consignment_docx(consignment, logo_path=r".\app\img\logo.jpg"):
    """
    Генерация документа DOCX с накладной.

    consignment: объект с атрибутами:
        - consignmentId: int
        - supplierName: str
        - date: str (YYYY-MM-DD)
        - items: list объектов с атрибутами name, quantity, price, sum
        - total: float
    logo_path: путь к файлу изображения логотипа
    """
    doc = Document()

    # Верхний блок: логотип слева + название компании справа
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    hdr_cells = header_table.rows[0].cells

    # Логотип
    if os.path.exists(logo_path):
        paragraph = hdr_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells[0].width = Inches(2)

    # Название компании
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.add_run("СушиСакура")
    run.bold = True
    run.font.size = Pt(32)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hdr_cells[1].width = Inches(4)

    doc.add_paragraph("\n")  # отступ после шапки

    # Заголовок накладной
    heading = doc.add_heading(level=0)
    run_heading = heading.add_run(f'Накладная №{consignment.consignmentId}')
    run_heading.font.size = Pt(24)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # Информация о поставщике
    info_paragraph = doc.add_paragraph()
    run = info_paragraph.add_run(f'Поставщик: {consignment.supplierName}\n')
    run.font.size = Pt(16)
    run = info_paragraph.add_run(f'Дата: {consignment.date}\n')
    run.font.size = Pt(16)
    doc.add_paragraph("\nТовары:\n").runs[0].font.size = Pt(16)

    # Таблица товаров
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Наименование', 'Кол-во', 'Цена', 'Сумма']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if len(hdr_cells[i].paragraphs[0].runs) > 0:
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(14)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Заполнение товаров
    for item in consignment.items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
        row_cells[1].text = f"{item.quantity:.2f}"
        row_cells[2].text = f"{item.price:.2f} ₽"
        row_cells[3].text = f"{item.sum:.2f} ₽"
        for cell in row_cells:
            if len(cell.paragraphs[0].runs) > 0:
                cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")  # отступ перед итогом

    # Итоговая сумма
    total_paragraph = doc.add_paragraph()
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_total = total_paragraph.add_run(f'Итого: {consignment.total:.2f} ₽')
    run_total.bold = True
    run_total.font.size = Pt(18)

    doc.add_paragraph("\n\n")  # отступ перед подписями

    # Строки для подписей
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'
    sig_cells = sig_table.rows[0].cells
    sig_cells[0].text = "Подпись ответственного: ______________________"
    sig_cells[1].text = "Подпись поставщика: ______________________"
    sig_cells = sig_table.rows[1].cells
    sig_cells[0].text = ""
    sig_cells[1].text = ""

    # Сохраняем на рабочий стол
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    filename = f"Накладная_{consignment.consignmentId}.docx"
    filepath = os.path.join(desktop, filename)
    doc.save(filepath)
    print(f"Накладная сохранена: {filepath}")

    return filepath
